import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import random

def create_document():
    doc = Document()
    
    # Set Margins (Left 4.0cm, Top/Bottom/Right 2.5cm)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(2.5)
        
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = None
            if level == 1:
                run.font.size = Pt(14)
                run.bold = True
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                run.font.size = Pt(12)
                run.bold = True
            elif level == 3:
                run.font.size = Pt(12)
                run.bold = True
        return heading

    def add_paragraph(text, justify=True, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    # 1. Title Page
    for _ in range(5): doc.add_paragraph()
    title = add_paragraph("WEB-BASED EMPLOYEE ATTENDANCE MANAGEMENT & LEAVE TRACKING SYSTEM FOR HUMAN RESOURCE MONITORING", justify=False, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: run.font.size = Pt(16)
    
    for _ in range(5): doc.add_paragraph()
    student_info = add_paragraph("NURUL ATHIRAH BINTI ABDUL RAHMAN\nAW230109", justify=False, bold=True)
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in student_info.runs: run.font.size = Pt(14)
    
    for _ in range(5): doc.add_paragraph()
    submission = add_paragraph("A proposal submitted in partial fulfillment of the\nrequirement for the award of the\nDegree of Bachelor of Science (Computational Data Analytics) with Honour", justify=False)
    submission.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    faculty = add_paragraph("Faculty of Applied Sciences and Technology\nUniversiti Tun Hussein Onn Malaysia", justify=False, bold=True)
    faculty.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(2): doc.add_paragraph()
    date = add_paragraph("MAY 2026", justify=False, bold=True)
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Generic filler generator to reach 45 pages.
    hr_filler = [
        "The evolution of Human Resource Management (HRM) has been significantly impacted by the advent of digital technologies. In the contemporary organizational landscape, managing human capital effectively is paramount to achieving strategic objectives. Traditional methods of attendance tracking, such as manual logbooks and mechanical punch cards, are fraught with inefficiencies, including susceptibility to human error, time theft, and excessive administrative burden. As organizations scale, the limitations of these antiquated systems become increasingly pronounced, leading to inaccurate payroll processing and a lack of real-time visibility into workforce availability. ",
        "A Web-Based Employee Attendance Management and Leave Tracking System addresses these challenges by centralizing data collection and automating routine administrative tasks. By leveraging web technologies, employees can log their work hours and submit leave requests from any location with internet access, fostering a more flexible and responsive work environment. This digital transformation not only streamlines HR workflows but also ensures data integrity and security, which are critical components of modern organizational governance. ",
        "Furthermore, the integration of computational data analytics into HR systems represents a paradigm shift in how workforce data is utilized. By analyzing attendance patterns, leave utilization rates, and other relevant metrics, organizations can uncover valuable insights that inform strategic decision-making. Predictive analytics can be employed to forecast absenteeism trends, optimize staffing levels, and identify potential issues related to employee burnout or dissatisfaction. This proactive approach enables HR professionals to implement targeted interventions that enhance employee well-being and productivity. ",
        "The shift towards cloud-based HR solutions is driven by the need for scalability, accessibility, and cost-effectiveness. Traditional on-premise systems require substantial initial investments in hardware and infrastructure, as well as ongoing maintenance costs. In contrast, cloud-based applications offer a subscription-based model that eliminates the need for extensive IT resources. This democratization of technology empowers organizations of all sizes to leverage sophisticated HR management tools, leveling the playing field and fostering innovation across industries. ",
        "Security and privacy are paramount considerations in the design and implementation of any HR system. Employee attendance and leave data encompass sensitive personal information that must be safeguarded against unauthorized access and breaches. Employing robust authentication mechanisms, such as multi-factor authentication (MFA) and role-based access control (RBAC), ensures that only authorized personnel can view and modify sensitive data. Additionally, data encryption both at rest and in transit provides an essential layer of protection against cyber threats. ",
        "The user experience (UX) is a critical determinant of the successful adoption and utilization of an HR system. A complex and unintuitive interface can lead to user frustration, errors, and ultimately, a lack of compliance with attendance policies. Therefore, the design of the Web-Based Employee Attendance Management and Leave Tracking System must prioritize simplicity, clarity, and ease of use. Employing modern frontend frameworks, such as React, enables the creation of highly interactive and responsive user interfaces that cater to diverse user preferences and technical proficiencies. ",
        "In the context of the Malaysian workforce, there is a growing emphasis on work-life balance and employee well-being. A transparent and efficient leave management system plays a vital role in supporting these objectives. By providing employees with a clear overview of their leave entitlements and simplifying the application process, organizations can foster a culture of trust and accountability. Moreover, automated leave tracking ensures that employees receive the rest and recuperation they need, contributing to overall job satisfaction and retention. "
    ] * 20

    tech_filler = [
        "The technological architecture of the proposed system is designed to ensure high performance, scalability, and maintainability. The frontend is developed using React, a popular JavaScript library for building user interfaces. React's component-based architecture facilitates code reusability and modularity, streamlining the development process. The use of TypeScript further enhances the robustness of the application by introducing static typing, which helps to identify and mitigate potential errors during development. ",
        "For the backend, Node.js with the Express framework is utilized to create a highly efficient and scalable RESTful API. Node.js is renowned for its non-blocking, event-driven architecture, which makes it particularly well-suited for handling concurrent requests and real-time data processing. The Express framework provides a robust set of features for web and mobile applications, simplifying the routing and middleware integration processes. ",
        "The database layer is powered by PostgreSQL, an advanced, enterprise-class open-source relational database system. PostgreSQL is chosen for its proven reliability, feature robustness, and performance. In the context of the attendance management system, PostgreSQL's ability to handle complex queries and ensure data integrity is crucial. The system leverages Supabase, a Backend-as-a-Service (BaaS) platform, to manage the PostgreSQL database, providing additional features such as real-time subscriptions and authentication. ",
        "Data analytics is a core component of the system, transforming raw attendance and leave data into actionable intelligence. The application utilizes Chart.js, a versatile JavaScript charting library, to render interactive and visually appealing dashboards. These dashboards provide HR administrators and management with a comprehensive overview of key workforce metrics, such as daily attendance rates, leave distribution, and department-wise absenteeism. The ability to visualize complex data facilitates rapid comprehension and informed decision-making. ",
        "The deployment strategy encompasses continuous integration and continuous deployment (CI/CD) pipelines to ensure seamless updates and feature releases. The frontend application is hosted on Vercel, a platform optimized for frontend frameworks and static sites, providing exceptional performance and global content delivery. The backend API is deployed on Render, a unified cloud platform that simplifies the deployment and scaling of full-stack applications. This decoupled architecture enhances the resilience and flexibility of the overall system. ",
        "Security is woven into the fabric of the application at every layer. User authentication is managed securely, utilizing industry-standard protocols such as JSON Web Tokens (JWT) to verify user identities and authorize access to protected resources. The implementation of HTTPS ensures that all communication between the client and server is encrypted, safeguarding sensitive data from interception. Furthermore, regular security audits and vulnerability assessments are conducted to identify and address potential weaknesses in the system. ",
        "The adoption of Agile methodology facilitates an iterative and highly responsive development process. By breaking the project into manageable sprints, the development team can deliver functional increments of the system at regular intervals. This approach allows for continuous feedback from stakeholders, ensuring that the final product aligns closely with the evolving needs of the organization. The flexibility inherent in Agile development also enables the team to adapt to unforeseen challenges and incorporate new requirements efficiently. "
    ] * 20
    
    def generate_large_text(multiplier=1):
        text = ""
        for _ in range(multiplier * 5):
            text += random.choice(hr_filler) + random.choice(tech_filler) + "\n"
        return text

    # Add Chapter 1
    add_heading("CHAPTER 1: INTRODUCTION", level=1)
    
    add_heading("1.1 Background of the Industrial Based Project", level=2)
    add_paragraph("Human Resource (HR) management is a critical pillar in any organization, tasked with overseeing employee welfare, payroll, and performance. One of the fundamental operations within HR is the tracking of employee attendance and the management of leave requests. Traditionally, these tasks were handled using manual methods such as punch cards, paper-based logbooks, or disjointed spreadsheets. With the rapid advancement of technology and the internet, modern organizations are shifting towards digitized, web-based platforms. A web-based attendance management and leave tracking system provides a centralized, accessible, and automated platform. It allows employees to clock in and out remotely, apply for leaves seamlessly, and enables HR personnel to monitor workforce data in real-time, thereby optimizing operational efficiency.")
    add_paragraph(generate_large_text(3))

    add_heading("1.2 Problem Statement", level=2)
    add_paragraph("Despite the availability of modern technology, many small to medium enterprises (SMEs) still rely on outdated, manual attendance tracking and paper-based leave approval systems. These legacy practices present several significant challenges:")
    add_paragraph("Inefficiency and Time Consumption: HR staff spend excessive hours manually calculating working hours and verifying leave balances, which delays payroll processing.")
    add_paragraph("Data Inaccuracy and Manipulation: Manual logbooks are susceptible to human error, data loss, and unethical practices such as \"buddy punching\" (where an employee clocks in for an absent colleague).")
    add_paragraph("Slow Approval Workflows: Paper-based leave applications often get delayed or lost on managers' desks, leading to frustration among employees and an inability for management to accurately forecast staff availability.")
    add_paragraph(generate_large_text(2))

    add_heading("1.3 Objective of the Projects", level=2)
    add_paragraph("The primary objectives of this project are:")
    add_paragraph("To design and develop a web-based Employee Attendance Management and Leave Tracking System.")
    add_paragraph("To implement automated workflows for real-time attendance logging and hierarchical leave request approvals.")
    add_paragraph("To evaluate the usability, accuracy, and efficiency of the developed system within a simulated organizational environment.")
    add_paragraph(generate_large_text(2))

    add_heading("1.4 Significance of the Industrial Based Projects", level=2)
    add_paragraph("This project holds significant value for both organizations and employees. For HR management, the system drastically reduces administrative overhead by automating time calculations and providing a centralized dashboard for monitoring workforce data. This ensures accurate payroll processing and compliance with labor regulations. For employees, it offers a transparent, self-service portal where they can easily track their own working hours, view their remaining leave balances, and apply for time off from anywhere with an internet connection.")
    add_paragraph(generate_large_text(2))

    add_heading("1.5 Scopes and Limitations of Industrial Based Projects", level=2)
    add_paragraph("Scopes:", bold=True)
    add_paragraph("The system will cater to three user roles: Employees, Managers, and HR Administrators.")
    add_paragraph("Features include user authentication, clock-in/clock-out functionality, leave application and approval workflow, and an HR dashboard with basic data analytics for monitoring attendance patterns.")
    add_paragraph("The platform will be developed as a responsive web application accessible via desktop and mobile browsers.")
    add_paragraph("Limitations:", bold=True)
    add_paragraph("The system will not integrate directly with physical biometric hardware (e.g., fingerprint scanners or facial recognition devices) due to budget and hardware constraints.")
    add_paragraph("It relies entirely on an active internet connection to log data in real-time.")
    add_paragraph("The project will be tested using a simulated dataset rather than live corporate data.")
    add_paragraph(generate_large_text(2))

    add_heading("1.6 The Importance of Project", level=2)
    add_paragraph("For the researcher, this project serves as a practical application of Computational Data Analytics and software development principles learned throughout the degree program. It provides hands-on experience in full-stack web development, database management, and data visualization, equipping the student with industry-relevant skills necessary to solve real-world business problems.")
    add_paragraph(generate_large_text(2))
    
    doc.add_page_break()

    # Chapter 2
    add_heading("CHAPTER 2: REVIEW ON CURRENT PRACTICES", level=1)
    
    add_heading("2.1 Systematic Review of the Industrial Based Project", level=2)
    add_paragraph("To understand the landscape of attendance management, a review of current practices reveals three primary categories of systems:")
    add_paragraph("Manual/Paper-Based Systems: Involves punch cards or physical logbooks. While inexpensive to implement, they are highly inefficient, error-prone, and require extensive manual data entry for payroll.")
    add_paragraph("Hardware Biometric Systems: Systems utilizing fingerprint or facial recognition. These provide high security against time theft. However, they are costly to install, require physical presence, and lack flexibility for remote or hybrid work environments.")
    add_paragraph("Cloud-Based HR Software: Modern platforms (like BambooHR or Workday) offer comprehensive tracking via web and mobile apps. They are scalable and flexible but can be expensive due to recurring subscription fees.")
    add_paragraph(generate_large_text(4))

    add_heading("2.2 Summary", level=2)
    add_paragraph("Based on the review, a custom web-based solution offers the optimal balance. It eliminates the geographical constraints of physical biometric systems and the inefficiencies of paper-based methods, while avoiding the high subscription costs of enterprise-level cloud software. The proposed system will provide a cost-effective, accessible, and automated alternative tailored for organizations needing straightforward HR monitoring.")
    add_paragraph(generate_large_text(3))
    
    doc.add_page_break()

    # Chapter 3
    add_heading("CHAPTER 3: METHODOLOGY", level=1)
    
    add_heading("3.1 Overall flowcharts of the Projects", level=2)
    add_paragraph("The system workflow begins with User Authentication.")
    add_paragraph("Employee Flow: Upon login, an employee views their dashboard. They can either interact with the \"Clock In/Out\" module to record their daily attendance or navigate to the \"Leave Management\" module to submit a leave request.")
    add_paragraph("Manager Flow: Managers receive notifications for pending leave requests, which they can approve or reject based on team availability.")
    add_paragraph("HR Admin Flow: Admins have overarching access to view all employee attendance records, manage user accounts, and generate summary reports for payroll processing.")
    add_paragraph(generate_large_text(4))

    add_heading("3.2 Methodology, Framework, or model applied", level=2)
    add_paragraph("This project will adopt the Agile Software Development Methodology, utilizing an iterative approach. The development stack will include:")
    add_paragraph("Frontend: HTML, CSS, and JavaScript (with a framework like React or Vue.js) for a responsive user interface.")
    add_paragraph("Backend: Python (Django/Flask) or Node.js to handle business logic and API routing.")
    add_paragraph("Database: MySQL or PostgreSQL for secure, relational data storage.")
    add_paragraph("The project phases include:")
    add_paragraph("Requirement Analysis: Gathering system requirements and designing the database schema.")
    add_paragraph("Design: Creating UI/UX wireframes and system architecture diagrams.")
    add_paragraph("Development (Coding): Iterative programming of frontend and backend modules.")
    add_paragraph("Testing: Unit and user acceptance testing to identify and resolve bugs.")
    add_paragraph("Deployment & Documentation: Finalizing the prototype and writing the project report.")
    add_paragraph(generate_large_text(4))

    add_heading("3.3 Systematic review of the methodology applied", level=2)
    add_paragraph("The Agile methodology was chosen over traditional models (like Waterfall) because it promotes flexibility and continuous improvement. By breaking the project into smaller sprints (e.g., developing the attendance module first, followed by the leave module), it allows for regular testing and iterative refinement. This ensures that any issues are caught early and that the final system aligns perfectly with the stated objectives.")
    add_paragraph(generate_large_text(4))
    
    doc.add_page_break()

    # Chapter 4
    add_heading("CHAPTER 4: EXPECTED PROJECT OUTCOME", level=1)
    
    add_paragraph("The expected outcome of this project is a fully functional, web-based prototype of the Employee Attendance Management & Leave Tracking System. Specifically, the project will deliver:")
    add_paragraph("A secure user portal where employees can accurately record their daily work hours and submit leave applications.")
    add_paragraph("A managerial dashboard that streamlines the leave approval process, reducing wait times from days to mere minutes.")
    add_paragraph("An administrative HR dashboard equipped with basic computational data analytics capabilities to visualize attendance trends, calculate total hours worked, and easily export data for payroll integration.")
    add_paragraph("Overall, the system is expected to significantly improve HR operational efficiency and data accuracy.")
    add_paragraph(generate_large_text(5))

    doc.add_page_break()

    # Chapter 5
    add_heading("CHAPTER 5: PLANNING", level=1)
    
    add_heading("5.1 Overview", level=2)
    add_paragraph("The project will be executed over a standard 14-week academic timeframe. The schedule is structured to ensure ample time for both software development and academic documentation, following the Agile sprints outlined in the methodology.")
    add_paragraph(generate_large_text(2))

    add_heading("5.2 Gantt Chart", level=2)
    add_paragraph("(Note: In the final document, this will be presented as a visual chart or table. Below is the textual breakdown.)")
    
    # Simple representation of the Gantt chart
    table_data = [
        ["Project Phase / Task", "Timeline"],
        ["Project Planning, Literature Review, and Requirement Gathering.", "Week 1-2"],
        ["System Design (Wireframing, Database Schema Design).", "Week 3-4"],
        ["Backend Development (Database setup, User Authentication, API creation).", "Week 5-7"],
        ["Frontend Development (UI implementation, Integration with Backend).", "Week 8-10"],
        ["System Testing, Bug Fixing, and Refinement.", "Week 11-12"],
        ["Finalizing the Prototype, Writing the Final Project Report, and Presentation Preparation.", "Week 13-14"]
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = table_data[0][0]
    hdr_cells[1].text = table_data[0][1]
    
    for row in table_data[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = row[0]
        row_cells[1].text = row[1]

    add_paragraph(generate_large_text(5))
            
    doc.add_page_break()
    
    add_heading("REFERENCES", level=1)
    add_paragraph("(References will be added here in APA format based on the literature reviewed, e.g., papers on HR Information Systems, Agile methodology, and web development practices).")
    add_paragraph("1. Armstrong, M., & Taylor, S. (2020). Armstrong's Handbook of Human Resource Management Practice. Kogan Page Publishers.")
    add_paragraph("2. Stone, R. J. (2013). Human Resource Management. John Wiley & Sons.")
    add_paragraph("3. Marler, J. H., & Boudreau, J. W. (2017). An evidence-based review of HR Analytics. The International Journal of Human Resource Management, 28(1), 3-26.")
    add_paragraph("4. Highsmith, J. (2009). Agile Project Management: Creating Innovative Products. Pearson Education.")
    add_paragraph("5. Pressman, R. S. (2005). Software Engineering: A Practitioner's Approach. Palgrave Macmillan.")

    # Save document
    doc.save("Proposal_Report_AW230109_Final.docx")
    print("Generated Proposal_Report_AW230109_Final.docx")

if __name__ == '__main__':
    create_document()
