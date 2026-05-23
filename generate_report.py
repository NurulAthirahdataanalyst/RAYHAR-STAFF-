import sys
import subprocess

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Heading styles
for i in range(1, 4):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = 'Times New Roman'
    h_font.size = Pt(12)
    h_font.bold = True
    h_font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    if i == 1:
        h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_style.paragraph_format.space_after = Pt(12)
        h_style.paragraph_format.space_before = Pt(24)

# Margin
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.left_margin = Cm(4.0)

# --- Title Page ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Web-Based Employee Attendance Management & Leave Tracking System for Human Resource Monitoring\n\n\n\nNURUL ATHIRAH BINTI ABDUL RAHMAN\nAW230109\n\n\n\n\nA proposal submitted in partial fulfillment of the requirement for the award of the\nDegree of Bachelor of Science (Computational Data Analytics) with Honour\n\n\n\n\nFaculty of Applied Sciences and Technology\nUniversiti Tun Hussein Onn Malaysia\n\n\n\n\nMAY 2026")
run.bold = True
doc.add_page_break()

# --- Content ---

# Chapter 1
doc.add_heading('CHAPTER 1\nINTRODUCTION', level=1)

doc.add_heading('1.1 Background of the Industrial Based Project', level=2)
doc.add_paragraph('Human Resource (HR) management is the backbone of organizational operations, responsible for managing employee welfare, tracking workforce productivity, and ensuring accurate payroll processing. One of the most fundamental yet labor-intensive tasks within HR is monitoring employee attendance and managing leave requests. Historically, organizations have relied on manual methods, such as paper-based logbooks, punch cards, or standalone spreadsheet files, to record when employees start and end their shifts.')
doc.add_paragraph('As digital transformation reshapes the modern workplace, there is a significant shift towards web-based HR monitoring systems. A Web-Based Employee Attendance Management and Leave Tracking System is a centralized digital platform that allows employees to clock in and out using web browsers and apply for leaves digitally. For HR administrators, it provides a real-time, consolidated view of workforce availability. By leveraging computational data analytics, these systems can process raw attendance timestamps to generate meaningful insights, such as identifying chronic absenteeism, tracking total hours worked, and automatically calculating remaining leave balances. This transition from manual to automated web-based tracking is crucial for enhancing operational efficiency, ensuring data integrity, and fostering a transparent work environment.')

doc.add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph('Despite the availability of software solutions, many organizations still utilize outdated, semi-manual processes for attendance and leave tracking. This project addresses several critical problems inherent in current legacy practices:')
doc.add_paragraph('1. Inefficient and Error-Prone Data Processing: HR personnel spend excessive time manually compiling attendance data from punch cards or scattered spreadsheets at the end of each month. This manual calculation of total hours worked and leave deductions frequently leads to human errors, resulting in inaccurate payroll processing and financial losses.')
doc.add_paragraph('2. Lack of Real-Time Monitoring and Data Integrity: Traditional systems are highly susceptible to time theft, specifically "buddy punching" (where an employee stamps a card for an absent colleague). Furthermore, management lacks a real-time overview of who is currently present, making daily workforce planning difficult.')
doc.add_paragraph('3. Cumbersome Leave Approval Workflows: Paper-based or email-based leave applications are inefficient. They are prone to being misplaced, delayed, or forgotten by managers. This creates frustration for employees who are unsure of their leave status and prevents HR from maintaining accurate, up-to-date leave balances.')
doc.add_paragraph('4. Absence of Data Analytics: Legacy systems store data statically, offering no analytical capabilities to identify trends such as high absenteeism rates in specific departments or seasonal leave patterns.')

doc.add_heading('1.3 Objective of the Projects', level=2)
doc.add_paragraph('The objectives of this project are structured to directly address the problem statement. The primary objectives are:')
doc.add_paragraph('1. To design and develop a secure, web-based Employee Attendance Management and Leave Tracking System tailored for Human Resource monitoring.')
doc.add_paragraph('2. To implement automated workflows that process real-time attendance timestamps and hierarchical leave approvals, integrating basic data analytics to generate visual HR reports.')
doc.add_paragraph('3. To evaluate the system\'s functionality, user interface (UI) usability, and data processing accuracy using simulated employee data within a controlled testing environment.')

doc.add_heading('1.4 Significance of the Industrial Based Projects', level=2)
doc.add_paragraph('The proposed system carries significant benefits for all stakeholders within an organization:')
doc.add_paragraph('For HR Administrators: It drastically reduces the administrative burden of calculating work hours and leave balances. The integration of data analytics allows HR to generate instant, accurate reports for payroll, saving days of manual labor each month.')
doc.add_paragraph('For Management: Managers gain a real-time dashboard displaying team availability, allowing for better daily resource allocation and project planning. The automated leave workflow ensures requests are processed promptly without email clutter.')
doc.add_paragraph('For Employees: It provides a transparent, self-service portal. Employees can easily view their attendance history, check their precise leave balances, and apply for time off from any device with an internet connection.')

doc.add_heading('1.5 Scopes and Limitations of Industrial Based Projects', level=2)
doc.add_paragraph('Scopes:')
doc.add_paragraph('Target Users: The system will feature three distinct user access levels: Employee, Manager, and HR Administrator.')
doc.add_paragraph('Core Modules: Attendance Module (web-based clock-in/clock-out functionality capturing server-side timestamps), Leave Management Module (digital application forms, automated routing to managers, and auto-updating of leave balances), and HR Monitoring Dashboard (visual data analytics showing daily attendance rates, absenteeism, and pending leave requests).')
doc.add_paragraph('Technology Stack: Development of a responsive web application utilizing HTML/CSS/JavaScript for the frontend, a robust backend framework (e.g., Node.js or Python Django), and a relational database (e.g., MySQL).')
doc.add_paragraph('Limitations:')
doc.add_paragraph('Hardware Integration: The system will operate purely as a software solution and will not integrate with physical biometric hardware (e.g., fingerprint scanners or facial recognition machines) or RFID card readers.')
doc.add_paragraph('Network Dependency: Real-time data logging and access to the system require a continuous, active internet connection. Offline attendance logging will not be supported in this prototype.')
doc.add_paragraph('Data Testing: The evaluation will be conducted using synthetic, simulated datasets rather than live, confidential corporate HR data.')

doc.add_heading('1.6 The Importance of Project', level=2)
doc.add_paragraph('For the researcher, this project represents a practical synthesis of the skills acquired in the Bachelor of Science in Computational Data Analytics program. It demonstrates the ability to not only develop functional software but to apply data analytics to solve tangible business problems. By structuring raw attendance data into visual HR insights, the project proves the student\'s competency in full-stack web development, database architecture, and data visualization.')

doc.add_page_break()

# Chapter 2
doc.add_heading('CHAPTER 2\nREVIEW ON CURRENT PRACTICES', level=1)

doc.add_heading('2.1 Systematic Review of the Industrial Based Project', level=2)
doc.add_paragraph('A systematic review of existing attendance and leave tracking methods highlights the strengths and weaknesses of current industry practices:')
doc.add_paragraph('1. Manual and Paper-Based Systems (Punch Cards & Logbooks):')
doc.add_paragraph('Overview: Employees physically stamp a card or sign a book upon entry and exit.')
doc.add_paragraph('Strengths: Extremely low initial setup cost; requires no technical training.')
doc.add_paragraph('Weaknesses: Highly vulnerable to buddy punching and data alteration. Calculating payroll requires painstakingly extracting data row by row, making it the most inefficient and error-prone method.')
doc.add_paragraph('2. Standalone Biometric Hardware Systems:')
doc.add_paragraph('Overview: Systems utilizing fingerprint scanners or facial recognition at office entrances.')
doc.add_paragraph('Strengths: Eradicates buddy punching entirely; highly secure.')
doc.add_paragraph('Weaknesses: High installation and maintenance costs. Data is often siloed in the machine and requires manual export via USB to HR computers. Furthermore, it completely fails to accommodate remote workers or employees out on field tasks.')
doc.add_paragraph('3. Enterprise Resource Planning (ERP) Cloud Software (e.g., Workday, SAP):')
doc.add_paragraph('Overview: Comprehensive, cloud-based HR management suites.')
doc.add_paragraph('Strengths: Highly scalable, accessible from anywhere, and features deep data analytics and automated workflows.')
doc.add_paragraph('Weaknesses: These systems are typically designed for massive corporations. They involve complex integration processes and carry prohibitive, recurring subscription fees that are unfeasible for SMEs.')

doc.add_heading('2.2 Summary', level=2)
doc.add_paragraph('The systematic review demonstrates a clear gap in the market for Small and Medium Enterprises (SMEs). Manual systems are too inefficient, biometric hardware is inflexible for modern remote work, and enterprise ERPs are too expensive. Therefore, the proposed custom Web-Based Attendance Management and Leave Tracking System bridges this gap. It provides the accessibility and automated analytics of an ERP without the prohibitive recurring costs, while eliminating the physical constraints of biometric hardware and the errors of manual logbooks.')

doc.add_page_break()

# Chapter 3
doc.add_heading('CHAPTER 3\nMETHODOLOGY', level=1)

doc.add_heading('3.1 Overall flowcharts of the Projects', level=2)
doc.add_paragraph('The system architecture and user flow are designed for logical progression:')
doc.add_paragraph('1. Authentication Flow: Users access the web portal and log in. The system authenticates credentials and routes the user to their specific dashboard (Employee, Manager, or Admin) based on role-based access control (RBAC).')
doc.add_paragraph('2. Attendance Flow (Employee): On the dashboard, the employee clicks the "Clock In" button. The system captures the current server time (preventing device-time manipulation) and logs the entry in the database. A "Clock Out" action completes the daily record.')
doc.add_paragraph('3. Leave Request Flow: An employee fills out a digital leave form (selecting dates and leave type). The system checks the database for sufficient leave balance. If valid, the status is set to "Pending" and a notification is triggered on the respective Manager\'s dashboard.')
doc.add_paragraph('4. Approval Flow (Manager): The manager reviews pending requests and selects "Approve" or "Reject". The database updates the leave status and recalculates the employee\'s remaining leave balance automatically.')
doc.add_paragraph('5. Monitoring Flow (HR Admin): The admin accesses an overarching dashboard where computational algorithms process the raw data to display total present staff, absent staff, and monthly attendance trends via visual charts.')

doc.add_heading('3.2 Methodology, Framework, or model applied', level=2)
doc.add_paragraph('This project will be developed using the Agile Software Development Methodology (Scrum Framework). Agile was selected because it allows for iterative development, breaking the project down into manageable sprints.')
doc.add_paragraph('Technology Framework:')
doc.add_paragraph('Frontend Presentation: HTML5, CSS3, and JavaScript, utilizing a library like Bootstrap or React to ensure the interface is mobile-responsive and user-friendly.')
doc.add_paragraph('Data Analytics & Visualization: Chart.js or D3.js will be used to render computational data analytics (e.g., pie charts for leave distribution, bar graphs for attendance trends).')
doc.add_paragraph('Backend Logic: Node.js (Express) or Python (Django) will manage the API endpoints, handle server-side timestamps, and execute the logic for leave balance deduction.')
doc.add_paragraph('Database Architecture: MySQL will be utilized as a relational database to maintain strict relationships between Users, Attendance Records, and Leave Requests tables.')

doc.add_heading('3.3 Systematic review of the methodology applied', level=2)
doc.add_paragraph('The traditional Waterfall model requires rigid, sequential phases, making it difficult to adapt to changes late in development. In contrast, the Agile methodology allows the system to be built incrementally. For example, Sprint 1 will focus purely on the Attendance module. Once that is tested and functional, Sprint 2 will focus on the Leave module. This iterative testing ensures that any logical flaws—such as incorrect leave deduction calculations—are identified and resolved immediately, ensuring a highly robust final product.')

doc.add_page_break()

# Chapter 4
doc.add_heading('CHAPTER 4\nEXPECTED PROJECT OUTCOME', level=1)
doc.add_paragraph('The successful execution of this project will result in a fully operational web application prototype. The specific expected outcomes include:')
doc.add_paragraph('1. A Functional Employee Portal: Allowing staff to seamlessly clock in/out and apply for leaves, eliminating the need for paper forms.')
doc.add_paragraph('2. An Automated Leave Management Engine: A backend system that automatically validates leave balances, routes requests to managers, and updates records upon approval without HR intervention.')
doc.add_paragraph('3. An HR Data Analytics Dashboard: A comprehensive monitoring screen that leverages computational data analytics to visually display workforce metrics. HR will be able to instantly view daily attendance percentages, track habitual tardiness, and export structured data sets for streamlined payroll processing.')
doc.add_paragraph('4. Project Documentation: A complete, academically rigorous report detailing the system architecture, database schema, and usability evaluation results.')

doc.add_page_break()

# Chapter 5
doc.add_heading('CHAPTER 5\nPLANNING', level=1)

doc.add_heading('5.1 Overview', level=2)
doc.add_paragraph('The project is planned over a standard 14-week academic semester. The schedule relies on the Agile methodology, ensuring continuous progression from requirement analysis through to final deployment and documentation.')

doc.add_heading('5.2 Gantt Chart Schedule', level=2)
doc.add_paragraph('The following represents the structured timeline of the project phases across 14 weeks:')
doc.add_paragraph('Phase 1: Research & Planning (Weeks 1-2). Includes Literature Review, Requirement Gathering, and Project Proposal Submission.')
doc.add_paragraph('Phase 2: System Design (Weeks 3-4). Includes UI/UX Wireframing, Database Schema, and Flowchart Design.')
doc.add_paragraph('Phase 3: Backend Development (Weeks 5-7). Includes Database Setup, User Authentication, and API Development for Attendance and Leave modules.')
doc.add_paragraph('Phase 4: Frontend & Analytics (Weeks 8-10). Includes UI Integration, Responsive Design, and Dashboard Data Analytics Integration.')
doc.add_paragraph('Phase 5: Testing & Refinement (Weeks 11-12). Includes Unit Testing, Bug Fixing, and Usability Evaluation with Simulated Data.')
doc.add_paragraph('Phase 6: Documentation (Weeks 13-14). Includes Final Report Writing and Final Presentation Preparation.')

doc.save(r'C:\Users\HP\ATTENDANCE_SYSTEM\Proposal_Report_AW230109.docx')
